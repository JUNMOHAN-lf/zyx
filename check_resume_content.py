import os
from docx import Document

# 检查Word文档的完整内容
def check_resume_content():
    """检查Word文档的完整内容"""
    file_path = "个人简历报告.docx"
    
    if not os.path.exists(file_path):
        print(f"文件不存在: {file_path}")
        return
    
    print(f"检查Word文档: {file_path}")
    print("=" * 80)
    
    try:
        doc = Document(file_path)
        
        print("\n=== 完整文档内容 ===")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                print(f"第{i+1}段: {text}")
        
        print("\n" + "=" * 80)
        print("=== 检查个人基本信息 ===")
        
        # 搜索可能的个人基本信息
        personal_info_keywords = ["姓名", "性别", "年龄", "专业", "学历", "联系方式", "电话", "邮箱", "地址"]
        found_info = False
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                for keyword in personal_info_keywords:
                    if keyword in text:
                        print(f"发现个人信息: {text}")
                        found_info = True
        
        if not found_info:
            print("未找到明确的个人基本信息")
            
    except Exception as e:
        print(f"读取文档失败: {e}")

if __name__ == "__main__":
    check_resume_content()
